"""艺康体 Word 导出：叙事结构 + 文件生成 + API 下载。"""

from __future__ import annotations

import importlib
import os
import sys
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from core import diagnostic as diag
from core import interactive as iv
from core import narrative as narr
from core import parser as pr
from core import report as report_mod
from data import make_sample

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if ROOT not in sys.path:
    sys.path.insert(0, ROOT)


@pytest.fixture
def sample_sess():
    data = pr.parse_financial_dict(make_sample.build_sample_data())
    diagnosis = diag.diagnose(data)
    sess = iv.start_session(data, diagnosis)
    while sess.state == iv.STATE_FINDING_LOOP and sess.current_finding:
        f = sess.current_finding
        iv.submit_decision(sess, f.id, "A")
    if sess.state == iv.STATE_CONFIRMATION:
        iv.confirm(sess, user_confirmed=True)
    return sess


def test_narrative_has_yikang_fields(sample_sess):
    n = narr.build_stage_narrative(
        sample_sess.data, sample_sess.diagnosis, sample_sess.decisions
    )
    assert n.one_liner
    assert n.year_rows
    assert n.now_points
    assert n.monthly_rows
    assert n.timeline
    assert n.methodology_notes
    assert n.subtitle.startswith("小白版")


def test_export_word_yikang_structure(sample_sess, tmp_path):
    path = tmp_path / "story.docx"
    report_mod.export_word(sample_sess, str(path))
    assert path.exists() and path.stat().st_size > 2000
    doc = Document(str(path))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    joined = "\n".join(texts)
    assert "经营业绩分析与建议" in joined or "示例制造" in joined
    assert "一、先说结论" in joined
    assert "五、用小白语言看懂关键指标" in joined
    assert "八、落地清单" in joined
    assert "九、数据口径说明" in joined
    assert len(doc.tables) >= 3  # 结论框/跨年/指标/月度/时间表


def test_export_api_word_download(sample_sess, tmp_path, monkeypatch):
    app_mod = importlib.import_module("web_backend.CO_app_WB-CO-TR-20260805160732")
    session = importlib.import_module("web_backend.CO_session_WB-CO-TR-20260805160732")
    interaction = importlib.import_module(
        "web_backend.CO_interaction_WB-CO-TR-20260805160732"
    )
    export_mod = importlib.import_module("web_backend.CO_export_WB-CO-TR-20260810")

    # 隔离导出目录
    out = tmp_path / "exports"
    out.mkdir()
    monkeypatch.setattr(export_mod, "_EXPORT_DIR", out)

    client = TestClient(app_mod.create_app())
    # 注入会话
    session.replace(
        sample_sess.data,
        ocr_texts=[],
        source_files=[],
        saved_previews=[],
    )
    with interaction._lock:
        interaction._sess = sample_sess
        interaction._sess_session_version = session.get_version()

    st = client.get("/api/export/status")
    assert st.status_code == 200
    body = st.json()
    assert body["unlocked"] is True

    r = client.post("/api/export/word")
    assert r.status_code == 200, r.text
    assert "wordprocessingml" in r.headers.get("content-type", "")
    assert len(r.content) > 2000
    # 落盘应存在
    files = list(out.glob("*.docx"))
    assert files, "应生成 docx 文件"
    client.close()
