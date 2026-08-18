"""利润宝 · Word/PDF 报告生成（S6）。

结构：封面 → 综合分析（指标表）→ 诊断发现 → 第二稿（决策表）→ 完整方案
（预计节税汇总）→ 合规声明 → 附录（增值税估算口径说明）。

- Word 用 python-docx；PDF 用 ReportLab，注册 CJK 字体避免中文乱码。
- 图表用 Matplotlib，自动探测中文字体（macOS PingFang SC / Linux Noto）。
- 导出失败抛 ReportError，调用方应保留内存方案并提示用户。
- 所有建议限于合法税务筹划；增值税税负率显著标注估算口径。
"""
from __future__ import annotations

import os
from datetime import datetime
from typing import List, Optional

from . import finance as fin
from . import industry as ind_mod
from .diagnostic import COMPLIANCE_NOTE, DiagnosisResult, Finding
from .interactive import Draft2Entry, Session
from .narrative import StageNarrative, build_stage_narrative


class ReportError(Exception):
    """报告生成错误。"""


# ── 中文字体探测（Matplotlib） ────────────────────────────────────────────
_CJK_FONT_CANDIDATES = [
    "PingFang SC", "Heiti SC", "STHeiti", "SimHei", "Microsoft YaHei",
    "Noto Sans CJK SC", "Noto Sans SC", "WenQuanYi Zen Hei", "Arial Unicode MS",
]
_resolved_font: Optional[str] = None


def _resolve_cjk_font() -> Optional[str]:
    """探测系统中文字体，返回可用字体名；无则返回 None。"""
    global _resolved_font
    if _resolved_font is not None:
        return _resolved_font
    try:
        import matplotlib
        matplotlib.use("Agg")  # 非交互后端
        from matplotlib import font_manager
        available = {f.name for f in font_manager.fontManager.ttflist}
        for name in _CJK_FONT_CANDIDATES:
            if name in available:
                _resolved_font = name
                return name
    except Exception:
        pass
    _resolved_font = ""  # 标记已尝试，避免重复探测
    return None or _resolved_font or None


def _setup_matplotlib_font():
    """配置 Matplotlib 中文字体；缺失时给出明确日志而非崩溃。"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        font = _resolve_cjk_font()
        if font:
            plt.rcParams["font.sans-serif"] = [font, "DejaVu Sans"]
            plt.rcParams["axes.unicode_minus"] = False
        return font
    except Exception:
        return None


def _fmt_money(v: float) -> str:
    return f"{v:,.0f} 元"


def _fmt_pct(v: float) -> str:
    return f"{v:.2f}%"


# ── 报告内容构建 ──────────────────────────────────────────────────────────

def _build_indicator_rows(sess: Session) -> List[List[str]]:
    """构造各年指标表行（年份 / 营收 / 增值税税负率 / 所得税税负率 / 毛利率 / 净利率）。"""
    rows: List[List[str]] = []
    for yr in sess.data.years:
        ind = fin.compute_year_indicators(sess.data, yr)
        vat = ind["增值税税负率"]
        rows.append([
            str(yr),
            _fmt_money(ind.get("营业收入", 0.0) or 0.0),
            f"{vat['value']}%（{vat['note']}）",
            _fmt_pct(ind["所得税税负率"]["value"]),
            _fmt_pct(ind["毛利率"]["value"]),
            _fmt_pct(ind["净利率"]["value"]),
        ])
    return rows


def _session_narrative(sess: Session) -> StageNarrative:
    """基于会话生成阶段叙事（含互动决策）。"""
    return build_stage_narrative(sess.data, sess.diagnosis, sess.decisions)


def _render_chart_to_png(sess: Session, tmp_dir: str) -> Optional[str]:
    """渲染营收与增值税税负率双轴图；返回 PNG 路径，失败返回 None。"""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        _setup_matplotlib_font()
        years = list(sess.data.years)
        if not years:
            return None
        revenues = []
        vat_rates = []
        for yr in years:
            ind = fin.compute_year_indicators(sess.data, yr)
            revenues.append(ind.get("营业收入", 0.0) or 0.0)
            vat_rates.append(ind["增值税税负率"]["value"])
        fig, ax1 = plt.subplots(figsize=(6, 3.2), dpi=120)
        color1 = "#2563eb"
        ax1.bar([str(y) for y in years], revenues, color=color1, alpha=0.7, label="营业收入")
        ax1.set_ylabel("营业收入（元）", color=color1)
        ax1.tick_params(axis="y", labelcolor=color1)
        ax2 = ax1.twinx()
        color2 = "#dc2626"
        ax2.plot([str(y) for y in years], vat_rates, color=color2, marker="o", linewidth=2, label="增值税税负率")
        ax2.set_ylabel("增值税税负率（%）", color=color2)
        ax2.tick_params(axis="y", labelcolor=color2)
        # 行业下限参考线
        benchmark, _ = ind_mod.get_benchmark(sess.data.industry)
        vat_min = benchmark["vat_tax_rate"]["min"]
        ax2.axhline(y=vat_min, color=color2, linestyle="--", alpha=0.4, label=f"行业下限 {vat_min:.1f}%")
        plt.title("营业收入与增值税税负率趋势（税负率为估算口径）")
        fig.tight_layout()
        path = os.path.join(tmp_dir, "trend.png")
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)
        return path
    except Exception:
        return None


# ── Word 报告 ─────────────────────────────────────────────────────────────

def _set_run_eastasia(run, font_name: str = "Songti SC") -> None:
    run.font.name = font_name
    try:
        from docx.oxml.ns import qn
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass


def _add_callout_table(doc, text: str, fill: str = "FFF3CD") -> None:
    """单格提示框（结论条 / 启发条），贴近艺康报告样式。"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    _set_run_eastasia(run)
    # 背景色
    try:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    except Exception:
        pass
    doc.add_paragraph("")


def export_word(
    sess: Session,
    path: str,
    ai_fallback: str = "",
    narrative: Optional[StageNarrative] = None,
) -> str:
    """生成艺康体小白版 Word 经营分析报告（.docx）。

    结构对齐《艺康装饰经营业绩分析与建议》：
    封面 → 一句话结论 → 先说结论 → 跨年表 → 最以前/中间/现在 →
    关键指标白话 → 将来建议 → 月度看板 → 落地清单 → 口径说明 →
    附录（诊断发现 + 互动决策）。

    narrative：可选的已合并叙事（DeepSeek 经营预算分析覆盖文本字段）；
    不传时用规则引擎叙事。表格/折线图等数字内容始终来自会话数据。
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        raise ReportError("未安装 python-docx，无法生成 Word 报告。") from e

    try:
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Songti SC"
        style.font.size = Pt(11)
        try:
            from docx.oxml.ns import qn
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        except Exception:
            pass

        narr = narrative if narrative is not None else _session_narrative(sess)
        years = list(sess.data.years or [])
        year_span = (
            f"{min(years)}-{max(years)}" if years else datetime.now().strftime("%Y")
        )

        # ── 封面（艺康体：公司名 + 年份 + 小白版副标题） ──
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(narr.company_name or sess.data.company_name or "未命名企业")
        r.bold = True
        r.font.size = Pt(18)
        _set_run_eastasia(r)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{year_span}经营业绩分析与建议")
        r.bold = True
        r.font.size = Pt(16)
        _set_run_eastasia(r)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(narr.subtitle)
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_run_eastasia(r)

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            (narr.data_source_note or "")
            + f"\n行业：{narr.industry}"
            + ("（基准已回退制造业）" if sess.diagnosis.industry_fallback else "")
        )
        r.font.size = Pt(10)
        _set_run_eastasia(r)

        # 一句话结论框
        if narr.one_liner:
            _add_callout_table(doc, narr.one_liner, fill="E8F4FD")

        # ── 一、先说结论 ──
        doc.add_heading("一、先说结论：公司这几年到底怎么样", level=1)
        # 把 headline 拆成多段（按句号）
        for part in (narr.headline or "").split("。"):
            part = part.strip()
            if part:
                doc.add_paragraph(part + "。")
        # 高风险发现再补 1-2 句
        highs = [f for f in sess.diagnosis.findings if f.severity == "高"][:2]
        for f in highs:
            doc.add_paragraph(f"{f.title}：{f.fact}")

        # 跨年对照表
        if narr.year_rows:
            doc.add_paragraph("")
            yt = doc.add_table(rows=1 + len(narr.year_rows), cols=6)
            yt.style = "Light Grid Accent 1"
            headers = ["年度", "收入", "净利润", "毛利率", "净利率", "一句话理解"]
            for i, h in enumerate(headers):
                yt.rows[0].cells[i].text = h
            for ri, row in enumerate(narr.year_rows):
                cells = yt.rows[ri + 1].cells
                cells[0].text = str(row.year)
                cells[1].text = row.revenue
                cells[2].text = row.net_profit
                cells[3].text = row.gross_margin
                cells[4].text = row.net_margin
                cells[5].text = row.one_liner

        # ── 二/三/四：阶段故事 ──
        stage_labels = []
        for st in narr.stages:
            if st.title.startswith("最以前"):
                stage_labels.append(("二", st))
            elif st.title.startswith("中间"):
                stage_labels.append(("三", st))
            else:
                stage_labels.append(("四", st))
        # 若只有 1-2 段，序号仍从二起
        if not stage_labels and narr.stages:
            stage_labels = [(str(i + 2), st) for i, st in enumerate(narr.stages)]

        for num, st in stage_labels:
            doc.add_heading(f"{num}、{st.title}", level=1)
            doc.add_paragraph(st.summary)
            for b in st.bullets:
                doc.add_paragraph(b, style="List Bullet")
            # 「现在」展开 H2 要点
            if st.title.startswith("现在") and narr.now_points:
                for pt in narr.now_points:
                    doc.add_heading(pt.title, level=2)
                    doc.add_paragraph(pt.body)

        if narr.stage_insight:
            _add_callout_table(doc, narr.stage_insight, fill="FFF8E7")

        # ── 五、关键指标白话 ──
        doc.add_heading("五、用小白语言看懂关键指标", level=1)
        if narr.now_metrics:
            mt = doc.add_table(rows=1 + len(narr.now_metrics), cols=4)
            mt.style = "Light Grid Accent 1"
            for i, h in enumerate(["指标", "小白解释", "当前情况", "管理判断"]):
                mt.rows[0].cells[i].text = h
            for ri, m in enumerate(narr.now_metrics):
                mt.rows[ri + 1].cells[0].text = m.name
                mt.rows[ri + 1].cells[1].text = m.plain
                mt.rows[ri + 1].cells[2].text = m.value_text
                mt.rows[ri + 1].cells[3].text = m.judgment

        # 税负估算注
        note_p = doc.add_paragraph()
        run = note_p.add_run(f"注：增值税税负率为 {fin.VAT_ESTIMATE_NOTE}，实际以申报数据为准。")
        run.bold = True
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        # 趋势图
        tmp_dir = os.path.dirname(path) or "."
        chart_path = _render_chart_to_png(sess, tmp_dir)
        if chart_path and os.path.exists(chart_path):
            try:
                doc.add_picture(chart_path, width=Cm(15))
                cap = doc.add_paragraph("图：营业收入与增值税税负率趋势")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception:
                pass

        # ── 六、将来 ──
        next_year = (max(years) + 1) if years else datetime.now().year + 1
        doc.add_heading(f"六、将来：{next_year}年以后建议怎么做", level=1)
        doc.add_paragraph(
            "未来不要只追求「收入做大」，更要追求「收入能赚钱、利润能回款、现金不断档」。"
            "建议按下面顺序推进。"
        )
        if narr.decision_summaries:
            doc.add_paragraph("（已结合互动决策）")
            for s in narr.decision_summaries:
                doc.add_paragraph(s, style="List Bullet")
        ordinals = ["第一", "第二", "第三", "第四", "第五", "第六", "第七", "第八", "第九", "第十"]
        for i, a in enumerate(narr.future_actions):
            if a.startswith("针对") or a[:1].isdigit():
                doc.add_paragraph(a)
            elif i < len(ordinals):
                doc.add_paragraph(f"{ordinals[i]}，{a}")
            else:
                doc.add_paragraph(f"{i + 1}. {a}")

        doc.add_paragraph(
            f"落地性评分：{sess.feasibility_score:.2f}%；"
            f"方案合计预计净影响：{_fmt_money(sess.total_est_saving)}。"
        )

        # ── 七、每月 8 个数 ──
        doc.add_heading("七、建议老板每月只盯这些数", level=1)
        if narr.monthly_rows:
            mt = doc.add_table(rows=1 + len(narr.monthly_rows), cols=3)
            mt.style = "Light Grid Accent 1"
            for i, h in enumerate(["每月看什么", "为什么看", "建议动作"]):
                mt.rows[0].cells[i].text = h
            for ri, row in enumerate(narr.monthly_rows):
                mt.rows[ri + 1].cells[0].text = row.metric
                mt.rows[ri + 1].cells[1].text = row.why
                mt.rows[ri + 1].cells[2].text = row.how
        else:
            for k in narr.monthly_kpis:
                doc.add_paragraph(k, style="List Bullet")

        # ── 八、落地清单 ──
        doc.add_heading("八、落地清单：从现在开始怎么动", level=1)
        if narr.timeline:
            mt = doc.add_table(rows=1 + len(narr.timeline), cols=3)
            mt.style = "Light Grid Accent 1"
            for i, h in enumerate(["时间", "要做的事", "目的"]):
                mt.rows[0].cells[i].text = h
            for ri, item in enumerate(narr.timeline):
                mt.rows[ri + 1].cells[0].text = item.when
                mt.rows[ri + 1].cells[1].text = item.action
                mt.rows[ri + 1].cells[2].text = item.purpose

        # ── 九、数据口径 ──
        doc.add_heading("九、数据口径说明", level=1)
        for note in narr.methodology_notes:
            doc.add_paragraph(note)
        doc.add_paragraph(
            "估算公式：估算增值税 = 税金及附加 ÷ 12%；税负率 = 估算增值税 ÷ 营业收入 × 100%。"
            "如能直连开票/申报系统可取真实应纳税额，则应替换为本口径。"
        )

        # ── 附录 A：诊断发现 ──
        doc.add_heading("附录 A：第一轮诊断发现", level=1)
        if sess.diagnosis.findings:
            doc.add_paragraph(
                f"共识别 {len(sess.diagnosis.findings)} 条问题/机会"
                f"（高 {sum(1 for f in sess.diagnosis.findings if f.severity=='高')} / "
                f"中 {sum(1 for f in sess.diagnosis.findings if f.severity=='中')} / "
                f"低 {sum(1 for f in sess.diagnosis.findings if f.severity=='低')}）。"
            )
            for f in sess.diagnosis.findings:
                doc.add_heading(f"{f.title}（严重度：{f.severity}）", level=2)
                doc.add_paragraph(f"事实：{f.fact}")
                doc.add_paragraph(f"行业对标：{f.benchmark}")
                doc.add_paragraph(f"初稿建议：{f.suggestion}")
                for opt in f.options:
                    doc.add_paragraph(
                        f"{opt.label}. {opt.name}：{opt.description}"
                        f"（目标 {opt.target_value:,.2f}{f.unit}；净影响 {_fmt_money(opt.est_saving)}）",
                        style="List Bullet",
                    )
        else:
            doc.add_paragraph("未发现明显异常，主要指标处于行业合理区间。")

        # ── 附录 B：互动第二稿 ──
        doc.add_heading("附录 B：互动决策与第二稿", level=1)
        if sess.draft2:
            t2 = doc.add_table(rows=1 + len(sess.draft2), cols=6)
            t2.style = "Light Grid Accent 1"
            for i, h in enumerate(["发现", "选项", "当前值", "目标值", "变动幅度", "预计净影响"]):
                t2.rows[0].cells[i].text = h
            for ri, e in enumerate(sess.draft2):
                cells = t2.rows[ri + 1].cells
                cells[0].text = e.finding_title
                cells[1].text = f"{e.option_label}. {e.option_name}"
                cells[2].text = f"{e.current_value:,.2f}"
                cells[3].text = f"{e.target_value:,.2f}"
                cells[4].text = e.change_pct
                cells[5].text = _fmt_money(e.est_saving)
            for e in sess.draft2:
                doc.add_heading(e.finding_title, level=3)
                doc.add_paragraph(f"选项：{e.option_label}. {e.option_name}")
                doc.add_paragraph(f"趋势：{e.trend}")
                doc.add_paragraph(f"操作细节：{e.action_detail}")
                doc.add_paragraph(f"注意事项：{e.cautions}")
        else:
            doc.add_paragraph("尚无互动决策记录。完成 A/B/C 互动后此处将展示第二稿明细。")

        if sess.strategy_notes:
            doc.add_heading("战略意图记录", level=2)
            for note in sess.strategy_notes:
                doc.add_paragraph(note, style="List Bullet")

        # ── 合规 ──
        doc.add_heading("合规声明", level=1)
        p = doc.add_paragraph(COMPLIANCE_NOTE)
        if p.runs:
            p.runs[0].bold = True
        doc.add_paragraph(
            "本报告优化建议限于合法税务筹划与经营管理范畴，包括但不限于：研发费用加计扣除、"
            "限额内据实扣除、业务模式与回款优化等。严禁虚开发票、隐匿收入、虚构成本。"
        )

        if ai_fallback:
            doc.add_heading("附录 C：AI 增强状态", level=1)
            doc.add_paragraph(ai_fallback)

        doc.save(path)
        return path
    except ReportError:
        raise
    except Exception as e:
        raise ReportError(f"Word 报告生成失败：{type(e).__name__}: {e}") from e


# ── PDF 报告（ReportLab + CJK 字体） ──────────────────────────────────────

# 系统中文字体候选路径（TrueType），用于 ReportLab 嵌入，避免 Poppler 下 CID 字体空白
_CJK_TTF_CANDIDATES = [
    # macOS
    "/System/Library/Fonts/PingFang.ttc",
    "/System/Library/Fonts/STHeiti Light.ttc",
    "/System/Library/Fonts/Hiragino Sans GB.ttc",
    "/Library/Fonts/Songti.ttc",
    # Linux (Noto / WenQuanYi)
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/wenquanyi/wqy-microhei/wqy-microhei.ttc",
    "/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc",
]


def _register_pdf_cjk_font():
    """注册 ReportLab CJK 字体；返回字体名。失败抛 ReportError。

    优先嵌入系统 TrueType 字体（PingFang/Noto）以确保跨查看器（含 Poppler）字形正确；
    若系统无 CJK TrueType，回退到 ReportLab 内置 STSong-Light CID 字体（部分查看器
    可能渲染异常，会在返回字体名后缀加 '_CID' 标记）。
    """
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except ImportError as e:
        raise ReportError("未安装 reportlab，无法生成 PDF 报告。") from e

    # 1) 优先尝试嵌入系统 TrueType 字体
    for ttf_path in _CJK_TTF_CANDIDATES:
        if not os.path.exists(ttf_path):
            continue
        try:
            # PingFang/Noto 是 ttc 集合，subfontIndex=0 取第一字重
            pdfmetrics.registerFont(TTFont("ProfitBaoCJK", ttf_path, subfontIndex=0))
            return "ProfitBaoCJK"
        except Exception:
            continue

    # 2) 回退到 ReportLab 内置 CID 字体（Poppler 下可能空白，但不会崩溃）
    font_name = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(font_name))
    except Exception:
        pass
    return font_name


def _clean_pdf_text(text: str) -> str:
    """清理 PDF 文本中的不兼容字符。

    - `&nbsp;` → 普通空格
    - `&amp;` / `&lt;` / `&gt;` → 实体转换
    - 项目符号 `•` → `*`（部分 CJK 字体无此字形）
    - 全角空格 `　` → 普通空格
    """
    if not text:
        return text
    return (
        text.replace("&nbsp;", " ")
            .replace("&amp;", "&")
            .replace("&lt;", "<")
            .replace("&gt;", ">")
            .replace("•", "*")
            .replace("　", " ")
    )


def export_pdf(
    sess: Session,
    path: str,
    ai_fallback: str = "",
    narrative: Optional[StageNarrative] = None,
) -> str:
    """生成 PDF 报告。返回路径；失败抛 ReportError。

    narrative：可选的已合并叙事（与 Word 同一份内容）；
    不传时用规则引擎叙事。指标表/折线图始终来自会话数据。
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak,
        )
    except ImportError as e:
        raise ReportError("未安装 reportlab，无法生成 PDF 报告。") from e

    cjk_font = _register_pdf_cjk_font()
    try:
        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=2 * cm, rightMargin=2 * cm,
            topMargin=2 * cm, bottomMargin=2 * cm,
            title="利润宝 · 企业财税优化方案",
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("H1CJK", parent=styles["Heading1"], fontName=cjk_font, fontSize=16, leading=22, textColor=colors.HexColor("#1e3a8a"))
        h2 = ParagraphStyle("H2CJK", parent=styles["Heading2"], fontName=cjk_font, fontSize=13, leading=18, textColor=colors.HexColor("#1e40af"))
        body = ParagraphStyle("BodyCJK", parent=styles["BodyText"], fontName=cjk_font, fontSize=10, leading=15)
        note_style = ParagraphStyle("NoteCJK", parent=body, textColor=colors.HexColor("#b91c1c"), fontName=cjk_font)
        title_style = ParagraphStyle("TitleCJK", parent=styles["Title"], fontName=cjk_font, fontSize=22, leading=30, alignment=1, textColor=colors.HexColor("#1e3a8a"))

        story = []
        narr = narrative if narrative is not None else _session_narrative(sess)
        # 封面
        story.append(Paragraph("利润宝 · 企业财税优化方案", title_style))
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(
            f"企业：{sess.data.company_name}<br/>行业：{sess.data.industry}<br/>"
            f"期间：{sess.data.years}<br/>生成日期：{datetime.now().strftime('%Y-%m-%d')}<br/>"
            "小白版：用看得懂的话，讲最以前、现在和将来",
            body,
        ))
        story.append(Spacer(1, 0.5 * cm))

        # 一、结论
        story.append(Paragraph("一、先说结论：公司这几年到底怎么样", h1))
        story.append(Paragraph(_clean_pdf_text(narr.headline), body))
        fallback_text = "（未匹配，已回退制造业基准）" if sess.diagnosis.industry_fallback else ""
        story.append(Paragraph(
            f"企业：{sess.data.company_name}；行业：{sess.data.industry}{fallback_text}；"
            f"年度：{', '.join(str(y) for y in sess.data.years)}",
            body,
        ))

        # 二、各阶段
        story.append(Paragraph("二、各阶段经营情况", h1))
        for st in narr.stages:
            story.append(Paragraph(_clean_pdf_text(st.title), h2))
            story.append(Paragraph(_clean_pdf_text(st.summary), body))
            for b in st.bullets:
                story.append(Paragraph(f"* {_clean_pdf_text(b)}", body))

        # 三、白话指标
        story.append(Paragraph("三、现在：关键指标白话解读", h1))
        if narr.now_metrics:
            m_data = [["指标", "小白解释", "当前情况", "管理判断"]]
            for m in narr.now_metrics:
                m_data.append([m.name, m.plain, m.value_text, m.judgment])
            mt = Table(m_data, repeatRows=1, colWidths=[3 * cm, 5 * cm, 3.5 * cm, 4 * cm])
            mt.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), cjk_font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(mt)
            story.append(Spacer(1, 0.3 * cm))

        # 四、指标表 + 图
        story.append(Paragraph("四、综合分析（年度指标表）", h1))
        rows = _build_indicator_rows(sess)
        table_data = [["年度", "营业收入", "增值税税负率", "所得税税负率", "毛利率", "净利率"]] + rows
        t = Table(table_data, repeatRows=1)
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), cjk_font),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ]))
        story.append(t)
        story.append(Spacer(1, 0.3 * cm))
        story.append(Paragraph(f"注：增值税税负率为 {fin.VAT_ESTIMATE_NOTE}，实际以申报数据为准。", note_style))
        story.append(Spacer(1, 0.3 * cm))

        tmp_dir = os.path.dirname(path) or "."
        chart_path = _render_chart_to_png(sess, tmp_dir)
        if chart_path and os.path.exists(chart_path):
            try:
                story.append(Image(chart_path, width=15 * cm, height=8 * cm))
                story.append(Paragraph("图：营业收入与增值税税负率趋势", body))
                story.append(Spacer(1, 0.3 * cm))
            except Exception:
                pass

        # 五、诊断发现
        story.append(PageBreak())
        story.append(Paragraph("五、第一轮诊断发现", h1))
        if sess.diagnosis.findings:
            story.append(Paragraph(
                f"共识别 {len(sess.diagnosis.findings)} 条问题/机会"
                f"（高 {sum(1 for f in sess.diagnosis.findings if f.severity=='高')} / "
                f"中 {sum(1 for f in sess.diagnosis.findings if f.severity=='中')} / "
                f"低 {sum(1 for f in sess.diagnosis.findings if f.severity=='低')}）。",
                body,
            ))
            for f in sess.diagnosis.findings:
                story.append(Paragraph(f"{f.title}（严重度：{f.severity}）", h2))
                story.append(Paragraph(f"事实：{_clean_pdf_text(f.fact)}", body))
                story.append(Paragraph(f"行业对标：{_clean_pdf_text(f.benchmark)}", body))
                story.append(Paragraph(f"初稿建议：{_clean_pdf_text(f.suggestion)}", body))
                story.append(Paragraph("可选方案：", body))
                for opt in f.options:
                    desc_text = _clean_pdf_text(opt.description)
                    story.append(Paragraph(
                        f"{opt.label}. {opt.name}<br/>"
                        f"  描述：{desc_text}<br/>"
                        f"  目标值：{opt.target_value:,.2f}；预计净影响：{_fmt_money(opt.est_saving)}；"
                        f"成本节约：{_fmt_money(opt.cost_saving)}；税收节约：{_fmt_money(opt.tax_saving)}；"
                        f"税负影响：{_fmt_money(opt.tax_impact)}；"
                        f"可行性：{opt.feasibility}；风险：{opt.risk_level}",
                        body,
                    ))
        else:
            story.append(Paragraph("未发现明显异常，主要指标处于行业合理区间。", body))

        # 六、第二稿
        story.append(Paragraph("六、第二稿（互动决策与增值测算）", h1))
        if sess.draft2:
            t2_data = [["发现", "选项", "当前值", "目标值", "变动幅度", "净影响"]]
            for e in sess.draft2:
                t2_data.append([
                    e.finding_title,
                    f"{e.option_label}. {e.option_name}",
                    f"{e.current_value:,.2f}",
                    f"{e.target_value:,.2f}",
                    e.change_pct,
                    _fmt_money(e.est_saving),
                ])
            t2 = Table(t2_data, repeatRows=1)
            t2.setStyle(TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), cjk_font),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e40af")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(t2)
            story.append(Spacer(1, 0.3 * cm))
            story.append(Paragraph("操作细节与注意事项", h2))
            for e in sess.draft2:
                story.append(Paragraph(e.finding_title, h2))
                story.append(Paragraph(f"选项：{e.option_label}. {e.option_name}", body))
                story.append(Paragraph(f"趋势（环比同比）：{_clean_pdf_text(e.trend)}", body))
                story.append(Paragraph(f"操作细节：{_clean_pdf_text(e.action_detail)}", body))
                story.append(Paragraph(f"注意事项：{_clean_pdf_text(e.cautions)}", body))
        else:
            story.append(Paragraph("无决策记录。完成互动问答后将展示决策明细。", body))

        # 七、将来
        story.append(Paragraph("七、将来要做什么（落地清单）", h1))
        if narr.decision_summaries:
            story.append(Paragraph("互动决策摘要", h2))
            for s in narr.decision_summaries:
                story.append(Paragraph(f"* {_clean_pdf_text(s)}", body))
        story.append(Paragraph(f"总净影响：{_fmt_money(sess.total_est_saving)}", body))
        story.append(Paragraph(f"落地性评分：{sess.feasibility_score:.2f}%", body))
        if sess.feasibility_breakdown:
            story.append(Paragraph("扣分明细：", body))
            for b in sess.feasibility_breakdown:
                story.append(Paragraph(f"- {_clean_pdf_text(b)}", body))
        for i, a in enumerate(narr.future_actions, 1):
            story.append(Paragraph(f"{i}. {_clean_pdf_text(a)}", body))
        story.append(Paragraph("建议老板每月盯住这些数", h2))
        for k in narr.monthly_kpis:
            story.append(Paragraph(f"* {_clean_pdf_text(k)}", body))
        if sess.strategy_notes:
            story.append(Paragraph("战略意图记录", h2))
            for n in sess.strategy_notes:
                story.append(Paragraph(f"- {_clean_pdf_text(n)}", body))

        # 八、合规
        story.append(Paragraph("八、合规声明", h1))
        story.append(Paragraph(COMPLIANCE_NOTE, ParagraphStyle("Bold", parent=body, fontName=cjk_font, textColor=colors.HexColor("#b91c1c"))))
        story.append(Paragraph(
            "本工具所有优化建议均属于合法税务筹划与经营管理建议范畴，包括但不限于：研发费用加计扣除、"
            "小微企业优惠、高新技术企业优惠、限额内费用据实扣除、业务模式与回款优化等。"
            "不构成投资、融资或法律意见，不替代审计报告。",
            body,
        ))

        # 九、口径
        story.append(Paragraph("九、增值税税负率估算口径说明", h1))
        story.append(Paragraph(
            "本报告所列增值税税负率为估算值，非真实应纳税额。"
            "估算公式：估算增值税 = 税金及附加 ÷ 12%（增值税附加税费占增值税比例的经验值），"
            "税负率 = 估算增值税 ÷ 营业收入 × 100%。"
            "如能直连开票/申报系统可取真实应纳税额，则应替换为本口径。"
            "实际税务申报与稽查以企业增值税申报表为准。",
            body,
        ))

        if ai_fallback:
            story.append(Paragraph("附录：AI 增强状态", h1))
            story.append(Paragraph(ai_fallback, body))

        doc.build(story)
        return path
    except ReportError:
        raise
    except Exception as e:
        raise ReportError(f"PDF 报告生成失败：{type(e).__name__}: {e}") from e


# ── T6.5 模板版预算报告（Word / PDF） ──────────────────────────────────────

def _budget_indicator_rows(plan) -> List[List[str]]:
    """构建经营目标指标行（顶部输入 + 计算结果）。"""
    ti = plan.top_inputs
    tc = plan.top_computed
    return [
        ["指标", "数值", "说明"],
        ["预算营业收入（C2）", f"{ti.budget_revenue:,.2f} 元", "顶部输入"],
        ["预算营业成本（C3）", f"{ti.budget_cost:,.2f} 元", "顶部输入"],
        ["毛利（C4=C2-C3）", f"{tc.gross_profit:,.2f} 元", "公式计算"],
        ["毛利率（C5=C4/C2）", f"{tc.gross_margin*100:.2f}%", "公式计算"],
        ["上年度营业收入（C6）", f"{ti.last_year_revenue:,.2f} 元", "顶部输入"],
        ["收入增长率（C7=(C2-C6)/C6）", f"{tc.revenue_growth_rate*100:.2f}%", "公式计算"],
        ["上年度营业成本（C8）", f"{ti.last_year_cost:,.2f} 元", "顶部输入"],
        ["上年度毛利率（C9=(C6-C8)/C6）", f"{tc.last_year_gross_margin*100:.2f}%", "公式计算"],
        ["行业所得税贡献率（E2）", f"{ti.industry_contribution_rate*100:.4f}%", "模板默认（待核验）"],
        ["企业预算所得税贡献率（E3）", f"{ti.company_contribution_rate*100:.4f}%", "顶部输入"],
        ["所得税税率（E4）", f"{ti.income_tax_rate*100:.2f}%", "默认高新 15%（可改选 5%/25%）"],
        ["应交所得税预算（E5=E3×C2）", f"{tc.income_tax_budget:,.2f} 元", "公式计算"],
        ["利润总额预算（E6=E5/E4）", f"{tc.profit_total_budget:,.2f} 元", "公式计算"],
        ["费用预算上限（E7=C4-E6）", f"{tc.expense_budget_cap:,.2f} 元", "公式计算"],
        ["实际已发生费用（E8=I98）", f"{tc.actual_expense_total:,.2f} 元", "明细联动"],
        ["费用差额（E9=E7-E8）", f"{tc.expense_diff:,.2f} 元", "正=剩余 负=超支"],
        ["未分配余额（G100=E7-G98）", f"{tc.unallocated_balance:,.2f} 元", "负=分配超上限"],
    ]


def _budget_summary_rows(plan) -> List[List[str]]:
    """构建预算执行汇总行。"""
    return [
        ["汇总项", "数值"],
        ["明细行数", f"{plan.total_lines} 行"],
        ["上年同期实际费用合计（D98）", f"{plan.last_year_total:,.2f} 元"],
        ["预算费用合计（G98）", f"{plan.allocated_total:,.2f} 元"],
        ["实际已发生费用合计（I98）", f"{plan.actual_total:,.2f} 元"],
        ["差额合计（J98）", f"{plan.diff_total:,.2f} 元"],
        ["超支项数", f"{plan.over_budget_count} 项"],
        ["临界项数", f"{plan.critical_count} 项"],
        ["待补录项数", f"{plan.pending_count} 项"],
        ["预算执行率", f"{(plan.actual_total/plan.allocated_total*100) if plan.allocated_total else 0:.1f}%"],
    ]


def export_budget_word(plan, path: str, session=None) -> str:
    """P0-4：生成模板版预算 Word 报告（含超支明细 + 诊断建议 + 负责人/期限 + 完整口径）。

    Args:
        plan: 预算计划
        path: 导出路径
        session: 可选诊断会话；有则写入诊断与优化建议章节，无则明确写「尚未完成诊断」

    Raises:
        ReportError: 校验未通过或生成失败；失败时不写入任何文件、不修改 plan
    """
    # P0-C（CO T7 重开）：导出前调用 validate_plan；失败时不写入文件、不修改 plan
    from .budget import validate_plan
    ok, errors = validate_plan(plan)
    if not ok:
        raise ReportError(
            "预算计划校验未通过，拒绝生成 Word：\n  - " + "\n  - ".join(errors)
            + "\n请先在「2. 模板工作台」修正输入后再导出。"
        )

    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError as e:
        raise ReportError("未安装 python-docx，无法生成 Word 报告。") from e

    try:
        doc = Document()
        # 默认中文字体
        style = doc.styles["Normal"]
        style.font.name = "Songti SC"
        style.font.size = Pt(10.5)
        try:
            style.element.rPr.rFonts.set(qn("w:eastAsia"), "Songti SC")
        except Exception:
            pass

        # 封面
        title = doc.add_heading("利润宝 · 经营目标与预算执行报告", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            f"企业：{plan.company_name or '（未命名）'}　　行业：{plan.industry}　　年度：{plan.year or '—'}",
        )
        doc.add_paragraph()

        # 一、经营目标
        doc.add_heading("一、经营目标（顶部公式联动）", level=1)
        _add_word_table(doc, _budget_indicator_rows(plan))
        doc.add_paragraph(
            "注：行业所得税贡献率与所得税税率为模板默认值，须以企业实际适用政策核验。",
        )

        # 二、预算执行汇总
        doc.add_heading("二、预算执行汇总", level=1)
        _add_word_table(doc, _budget_summary_rows(plan))

        # 三、超支异常明细（P0-4 新增）
        doc.add_heading("三、超支异常明细", level=1)
        over_lines = [l for l in plan.lines if l.exec_status == "超支"]
        critical_lines = [l for l in plan.lines if l.exec_status == "临界"]
        doc.add_paragraph(
            f"超支项数：{plan.over_budget_count}　临界项数：{plan.critical_count}　待补录项数：{plan.pending_count}",
        )
        if over_lines:
            over_header = ["行号", "科目", "费用名称", "预算费用", "实际已发生", "差额", "执行率", "执行状态"]
            over_rows = [over_header]
            for l in over_lines:
                exec_rate = (l.actual_amount / l.budget_amount * 100) if l.budget_amount else 0
                over_rows.append([
                    str(l.row), l.subject, l.expense_name,
                    f"{l.budget_amount:,.0f}", f"{l.actual_amount:,.0f}",
                    f"{l.diff:,.0f}", f"{exec_rate:.1f}%", l.exec_status,
                ])
            _add_word_table(doc, over_rows)
            doc.add_paragraph(
                "处理建议：超支项应优先复核业务真实性与必要性；可在「3. 互动」环节通过 A/B/C 选项生成压降方案。",
            )
        else:
            doc.add_paragraph("无超支项。")

        # 四、费用明细（84 行）
        doc.add_heading("四、费用明细（84 行）", level=1)
        detail_header = ["行号", "科目", "费用名称", "上年实际", "预算费用", "实际已发生", "差额", "执行状态"]
        detail_rows = [detail_header]
        for l in plan.lines:
            detail_rows.append([
                str(l.row), l.subject, l.expense_name,
                f"{l.last_year_actual:,.0f}", f"{l.budget_amount:,.0f}",
                f"{l.actual_amount:,.0f}", f"{l.diff:,.0f}", l.exec_status,
            ])
        _add_word_table(doc, detail_rows)

        # 五、诊断与优化建议（P0-4 新增；含负责人/期限/成本节约/税收节约/税负影响/净影响分栏）
        doc.add_heading("五、诊断与优化建议", level=1)
        if session is not None and getattr(session, "draft2", None):
            doc.add_paragraph(
                f"总成本节约：{_fmt_money(sum(getattr(d, 'cost_saving', 0.0) or 0.0 for d in session.draft2))}　"
                f"总税收节约：{_fmt_money(sum(getattr(d, 'tax_saving', 0.0) or 0.0 for d in session.draft2))}　"
                f"总税负影响：{_fmt_money(sum(getattr(d, 'tax_impact', 0.0) or 0.0 for d in session.draft2))}　"
                # P0-B（CO T7 重开）：Draft2Entry 字段为 est_saving（非 net_impact）
                f"总净影响：{_fmt_money(sum(getattr(d, 'est_saving', 0.0) or 0.0 for d in session.draft2))}",
            )
            action_header = [
                "序号", "发现", "选项", "目标值",
                "成本节约(元)", "税收节约(元)", "税负影响(元)", "净影响(元)",
                "负责人", "期限", "执行状态",
            ]
            action_rows = [action_header]
            for idx, d in enumerate(session.draft2, start=1):
                action_rows.append([
                    str(idx),
                    getattr(d, "finding_title", "") or getattr(d, "finding_id", ""),
                    f"{getattr(d, 'option_label', '')}. {getattr(d, 'option_name', '')}",
                    f"{getattr(d, 'target_value', 0):,.2f}",
                    f"{getattr(d, 'cost_saving', 0.0) or 0.0:,.0f}",
                    f"{getattr(d, 'tax_saving', 0.0) or 0.0:,.0f}",
                    f"{getattr(d, 'tax_impact', 0.0) or 0.0:,.0f}",
                    # P0-B：Draft2Entry 字段为 est_saving
                    f"{getattr(d, 'est_saving', 0.0) or 0.0:,.0f}",
                    getattr(d, "owner", "") or "（待指派）",
                    getattr(d, "deadline", "") or "（待确定）",
                    getattr(d, "exec_status", "待执行") if hasattr(d, "exec_status") else "待执行",
                ])
            _add_word_table(doc, action_rows)
            # 操作细节与注意事项
            doc.add_heading("操作细节与注意事项", level=2)
            for d in session.draft2:
                doc.add_heading(getattr(d, "finding_title", ""), level=3)
                doc.add_paragraph(f"选项：{getattr(d, 'option_label', '')}. {getattr(d, 'option_name', '')}")
                doc.add_paragraph(f"目标值：{getattr(d, 'target_value', 0):,.2f}")
                if hasattr(d, "trend"):
                    doc.add_paragraph(f"趋势（环比同比）：{d.trend}")
                if hasattr(d, "action_detail"):
                    doc.add_paragraph(f"操作细节：{d.action_detail}")
                if hasattr(d, "cautions"):
                    doc.add_paragraph(f"注意事项：{d.cautions}")
        else:
            doc.add_paragraph(
                "尚未完成诊断；请先在「1. 导入」载入三年财报 → 「2. 诊断」执行第一轮 → "
                "「3. 互动」完成 A/B/C 决策，再生成本章节的优化建议。本工具不得伪造建议。",
            )

        # 六、完整计算口径（P0-4 新增）
        doc.add_heading("六、完整计算口径", level=1)
        doc.add_paragraph(
            "顶部公式：\n"
            "  C4 毛利 = C2 预算营业收入 − C3 预算营业成本\n"
            "  C5 毛利率 = IF(C2=0, 0, C4 / C2)\n"
            "  C7 收入增长率 = IF(C6=0, 0, (C2 − C6) / C6)\n"
            "  C9 上年度毛利率 = IF(C6=0, 0, (C6 − C8) / C6)\n"
            "  E5 应交所得税预算 = E3 企业所得税贡献率 × C2\n"
            "  E6 利润总额预算 = IF(E4=0, 0, E5 / E4)\n"
            "  E7 费用预算上限 = C4 − E6\n"
            "  E8 实际已发生费用 = I98（明细 I 列合计）\n"
            "  E9 费用差额 = E7 − E8（正=剩余，负=超支）\n"
            "  G100 未分配余额 = E7 − G98\n"
            "明细公式：\n"
            "  E 上年费用率 = IF($C$6=0, 0, D / $C$6)\n"
            "  F 参考金额 = D × (1 + $C$7)\n"
            "  H 预算费用率 = IF($C$2=0, 0, G / $C$2)\n"
            "  J 差额 = G − I\n"
            "执行状态判定：待补录（G=I=0）/ 正常（执行率<80%）/ 临界（80%-100%）/ 超支（>100%）",
        )

        # 七、合规声明
        doc.add_heading("七、合规声明", level=1)
        doc.add_paragraph(COMPLIANCE_NOTE)
        doc.add_paragraph(
            "增值税税负率为 " + fin.VAT_ESTIMATE_NOTE
            + "；公式：税金及附加 ÷ 12% ÷ 营业收入 × 100%。实际以申报数据为准。",
        )

        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc.save(path)
        return path
    except ReportError:
        raise
    except Exception as e:
        raise ReportError(f"预算 Word 报告生成失败：{type(e).__name__}: {e}") from e


def _add_word_table(doc, rows):
    """向 Word 文档添加表格（首行为表头）。"""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            if i == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True


def export_budget_pdf(plan, path: str, session=None) -> str:
    """P0-4：生成模板版预算 PDF 报告（含超支明细 + 诊断建议 + 负责人/期限 + 完整口径）。

    Args:
        plan: 预算计划
        path: 导出路径
        session: 可选诊断会话；有则写入诊断与优化建议章节，无则明确写「尚未完成诊断」

    Raises:
        ReportError: 校验未通过或生成失败；失败时不写入任何文件、不修改 plan
    """
    # P0-C（CO T7 重开）：导出前调用 validate_plan；失败时不写入文件、不修改 plan
    from .budget import validate_plan
    ok, errors = validate_plan(plan)
    if not ok:
        raise ReportError(
            "预算计划校验未通过，拒绝生成 PDF：\n  - " + "\n  - ".join(errors)
            + "\n请先在「2. 模板工作台」修正输入后再导出。"
        )

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
        )
    except ImportError as e:
        raise ReportError("未安装 reportlab，无法生成 PDF 报告。") from e

    cjk_font = _register_pdf_cjk_font()
    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        doc = SimpleDocTemplate(path, pagesize=A4,
                                leftMargin=1.8*cm, rightMargin=1.8*cm,
                                topMargin=1.8*cm, bottomMargin=1.8*cm)
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("H1", parent=styles["Heading1"],
                            fontName=cjk_font, fontSize=15, leading=20)
        h2 = ParagraphStyle("H2", parent=styles["Heading2"],
                            fontName=cjk_font, fontSize=12, leading=16)
        h3 = ParagraphStyle("H3", parent=styles["Heading3"],
                            fontName=cjk_font, fontSize=10.5, leading=14)
        body = ParagraphStyle("Body", parent=styles["Normal"],
                              fontName=cjk_font, fontSize=9.5, leading=14)
        note = ParagraphStyle("Note", parent=body, textColor=colors.red, fontSize=8.5)

        story = []
        # 封面
        story.append(Paragraph("利润宝 · 经营目标与预算执行报告", h1))
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            _clean_pdf_text(f"企业：{plan.company_name or '（未命名）'}　行业：{plan.industry}　年度：{plan.year or '—'}"),
            body,
        ))
        story.append(Spacer(1, 12))

        _pdf_table_style = TableStyle([
            ("FONT", (0, 0), (-1, -1), cjk_font, 8),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ])

        # 一、经营目标
        story.append(Paragraph("一、经营目标（顶部公式联动）", h2))
        ind_rows = _budget_indicator_rows(plan)
        ind_data = [[_clean_pdf_text(c) for c in row] for row in ind_rows]
        story.append(Table(ind_data, repeatRows=1, style=_pdf_table_style))
        story.append(Spacer(1, 4))
        story.append(Paragraph(
            _clean_pdf_text("注：行业所得税贡献率与所得税税率为模板默认值，须以企业实际适用政策核验。"),
            note,
        ))
        story.append(Spacer(1, 10))

        # 二、预算执行汇总
        story.append(Paragraph("二、预算执行汇总", h2))
        sum_rows = _budget_summary_rows(plan)
        sum_data = [[_clean_pdf_text(c) for c in row] for row in sum_rows]
        story.append(Table(sum_data, repeatRows=1, style=_pdf_table_style))
        story.append(Spacer(1, 10))

        # 三、超支异常明细（P0-4 新增）
        story.append(Paragraph("三、超支异常明细", h2))
        over_lines = [l for l in plan.lines if l.exec_status == "超支"]
        story.append(Paragraph(
            _clean_pdf_text(f"超支项数：{plan.over_budget_count}　临界项数：{plan.critical_count}　待补录项数：{plan.pending_count}"),
            body,
        ))
        if over_lines:
            over_header = ["行号", "科目", "费用名称", "预算费用", "实际已发生", "差额", "执行率", "执行状态"]
            over_data = [over_header]
            for l in over_lines:
                exec_rate = (l.actual_amount / l.budget_amount * 100) if l.budget_amount else 0
                over_data.append([
                    str(l.row), _clean_pdf_text(l.subject),
                    _clean_pdf_text(l.expense_name),
                    f"{l.budget_amount:,.0f}", f"{l.actual_amount:,.0f}",
                    f"{l.diff:,.0f}", f"{exec_rate:.1f}%",
                    _clean_pdf_text(l.exec_status),
                ])
            story.append(Table(over_data, repeatRows=1, style=_pdf_table_style))
            story.append(Spacer(1, 4))
            story.append(Paragraph(
                _clean_pdf_text("处理建议：超支项应优先复核业务真实性与必要性；可在「3. 互动」环节通过 A/B/C 选项生成压降方案。"),
                body,
            ))
        else:
            story.append(Paragraph("无超支项。", body))
        story.append(Spacer(1, 10))

        # 四、费用明细（84 行）
        story.append(Paragraph("四、费用明细（84 行）", h2))
        detail_header = ["行号", "科目", "费用名称", "上年实际", "预算", "实际", "差额", "状态"]
        detail_data = [detail_header]
        for l in plan.lines:
            detail_data.append([
                str(l.row), _clean_pdf_text(l.subject),
                _clean_pdf_text(l.expense_name),
                f"{l.last_year_actual:,.0f}", f"{l.budget_amount:,.0f}",
                f"{l.actual_amount:,.0f}", f"{l.diff:,.0f}",
                _clean_pdf_text(l.exec_status),
            ])
        story.append(Table(detail_data, repeatRows=1, style=TableStyle([
            ("FONT", (0, 0), (-1, -1), cjk_font, 7.5),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
            ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ])))
        story.append(Spacer(1, 10))

        # 五、诊断与优化建议（P0-4 新增）
        story.append(Paragraph("五、诊断与优化建议", h2))
        if session is not None and getattr(session, "draft2", None):
            total_cost = sum(getattr(d, "cost_saving", 0.0) or 0.0 for d in session.draft2)
            total_tax = sum(getattr(d, "tax_saving", 0.0) or 0.0 for d in session.draft2)
            total_impact = sum(getattr(d, "tax_impact", 0.0) or 0.0 for d in session.draft2)
            total_net = sum(getattr(d, "est_saving", 0.0) or 0.0 for d in session.draft2)  # P0-B：est_saving 非 net_impact
            story.append(Paragraph(
                _clean_pdf_text(f"总成本节约：{total_cost:,.0f} 元　总税收节约：{total_tax:,.0f} 元　"
                                f"总税负影响：{total_impact:,.0f} 元　总净影响：{total_net:,.0f} 元"),
                body,
            ))
            story.append(Spacer(1, 4))
            action_header = [
                "序号", "发现", "选项", "目标值",
                "成本节约", "税收节约", "税负影响", "净影响",
                "负责人", "期限", "执行状态",
            ]
            action_data = [action_header]
            for idx, d in enumerate(session.draft2, start=1):
                action_data.append([
                    str(idx),
                    _clean_pdf_text(getattr(d, "finding_title", "") or getattr(d, "finding_id", "")),
                    _clean_pdf_text(f"{getattr(d, 'option_label', '')}. {getattr(d, 'option_name', '')}"),
                    f"{getattr(d, 'target_value', 0):,.0f}",
                    f"{getattr(d, 'cost_saving', 0.0) or 0.0:,.0f}",
                    f"{getattr(d, 'tax_saving', 0.0) or 0.0:,.0f}",
                    f"{getattr(d, 'tax_impact', 0.0) or 0.0:,.0f}",
                    # P0-B：Draft2Entry 字段为 est_saving
                    f"{getattr(d, 'est_saving', 0.0) or 0.0:,.0f}",
                    _clean_pdf_text(getattr(d, "owner", "") or "（待指派）"),
                    _clean_pdf_text(getattr(d, "deadline", "") or "（待确定）"),
                    _clean_pdf_text(getattr(d, "exec_status", "待执行") if hasattr(d, "exec_status") else "待执行"),
                ])
            story.append(Table(action_data, repeatRows=1, style=TableStyle([
                ("FONT", (0, 0), (-1, -1), cjk_font, 7),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
                ("FONTSIZE", (0, 0), (-1, -1), 7),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ])))
            story.append(Spacer(1, 6))
            story.append(Paragraph("操作细节与注意事项", h3))
            for d in session.draft2:
                story.append(Paragraph(
                    _clean_pdf_text(f"· {getattr(d, 'finding_title', '')}："
                                    f"选项 {getattr(d, 'option_label', '')}. {getattr(d, 'option_name', '')}；"
                                    f"目标值 {getattr(d, 'target_value', 0):,.2f}"),
                    body,
                ))
                if hasattr(d, "action_detail"):
                    story.append(Paragraph(_clean_pdf_text(f"  操作细节：{d.action_detail}"), body))
                if hasattr(d, "cautions"):
                    story.append(Paragraph(_clean_pdf_text(f"  注意事项：{d.cautions}"), body))
        else:
            story.append(Paragraph(
                _clean_pdf_text("尚未完成诊断；请先在「1. 导入」载入三年财报 → 「2. 诊断」执行第一轮 → "
                                "「3. 互动」完成 A/B/C 决策，再生成本章节的优化建议。本工具不得伪造建议。"),
                body,
            ))
        story.append(Spacer(1, 10))

        # 六、完整计算口径（P0-4 新增）
        story.append(Paragraph("六、完整计算口径", h2))
        formula_text = (
            "顶部公式：<br/>"
            "  C4 毛利 = C2 预算营业收入 - C3 预算营业成本<br/>"
            "  C5 毛利率 = IF(C2=0, 0, C4 / C2)<br/>"
            "  C7 收入增长率 = IF(C6=0, 0, (C2 - C6) / C6)<br/>"
            "  C9 上年度毛利率 = IF(C6=0, 0, (C6 - C8) / C6)<br/>"
            "  E5 应交所得税预算 = E3 企业所得税贡献率 × C2<br/>"
            "  E6 利润总额预算 = IF(E4=0, 0, E5 / E4)<br/>"
            "  E7 费用预算上限 = C4 - E6<br/>"
            "  E8 实际已发生费用 = I98（明细 I 列合计）<br/>"
            "  E9 费用差额 = E7 - E8（正=剩余，负=超支）<br/>"
            "  G100 未分配余额 = E7 - G98<br/>"
            "明细公式：<br/>"
            "  E 上年费用率 = IF($C$6=0, 0, D / $C$6)<br/>"
            "  F 参考金额 = D × (1 + $C$7)<br/>"
            "  H 预算费用率 = IF($C$2=0, 0, G / $C$2)<br/>"
            "  J 差额 = G - I<br/>"
            "执行状态判定：待补录（G=I=0）/ 正常（执行率&lt;80%）/ 临界（80%-100%）/ 超支（&gt;100%）"
        )
        story.append(Paragraph(formula_text, body))
        story.append(Spacer(1, 10))

        # 七、合规声明
        story.append(Paragraph("七、合规声明", h2))
        story.append(Paragraph(_clean_pdf_text(COMPLIANCE_NOTE), body))
        story.append(Paragraph(
            _clean_pdf_text("增值税税负率为 " + fin.VAT_ESTIMATE_NOTE
                            + "；实际以申报数据为准。"),
            note,
        ))

        doc.build(story)
        return path
    except ReportError:
        raise
    except Exception as e:
        raise ReportError(f"预算 PDF 报告生成失败：{type(e).__name__}: {e}") from e
